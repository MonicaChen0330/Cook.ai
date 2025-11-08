from . import Document, Page, ExtractedImage, DocumentLoader
from docx import Document as DocxDocument
from .image_utils import image_to_base64_uri
from .ocr_utils import ocr_image_to_text

class DocxLoader(DocumentLoader):
    """A loader for Microsoft Word (.docx) files."""

    def load(self, source: str) -> Document:
        """Reads text and extracts images from a .docx file."""
        try:
            doc = DocxDocument(source)
            native_text_parts = []
            doc_images = []

            for para in doc.paragraphs:
                native_text_parts.append(para.text)
            
            # Extract images from inline shapes
            for rel in doc.part.rels:
                if "image" in doc.part.rels[rel].target_ref:
                    try:
                        image_part = doc.part.rels[rel].target_part
                        image_bytes = image_part.blob
                        
                        ocr_text = ocr_image_to_text(image_bytes) or ""
                        base64_image_uri = image_to_base64_uri(image_bytes)

                        if base64_image_uri:
                            doc_images.append(ExtractedImage(image_uri=base64_image_uri, ocr_text=ocr_text))
                    except Exception as e:
                        print(f"Warning: Could not process an image in {source}: {e}")

            # Consolidate all text into a single string for one page
            native_text = "\n".join(native_text_parts)
            single_page = Page(page_number=1, native_text=native_text, extracted_images=doc_images)

            print(f"Successfully read content and {len(doc_images)} images from {source}")
            return Document(source=source, pages=[single_page])
        except Exception as e:
            print(f"Error reading DOCX file: {str(e)}")
            raise e
